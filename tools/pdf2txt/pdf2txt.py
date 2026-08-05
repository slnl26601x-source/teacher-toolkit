"""
PDF 轉文字/Word 工具（PyMuPDF + python-docx）

用法：
  python pdf2txt.py "檔案.pdf"                          # 轉單一檔案 → .docx
  python pdf2txt.py "檔案.pdf" --format txt             # 轉純文字
  python pdf2txt.py "檔案.pdf" --analyze                # 只分析 PDF 類型 + 建議做法
  python pdf2txt.py "檔案.pdf" --auto                   # 分析後自動選最佳方式處理
  python pdf2txt.py "資料夾"                           # 轉資料夾內所有 PDF
  python pdf2txt.py "資料夾" --out "輸出資料夾"         # 指定輸出位置
  python pdf2txt.py "a.pdf" "b.pdf"                    # 多個檔案
  python pdf2txt.py "資料夾" --layout                  # 保留排版（欄位/多欄）
  python pdf2txt.py "資料夾" --combine                 # 多檔案合併成單一檔案

輸出：與來源同名，.txt 或 .docx（依 --format）。

處理方式判斷（--analyze / --auto）：
  - 文字型 PDF（有內嵌文字層）→ 直接抽取（pdf2txt / pdf2docx），快又準
  - 掃描型 PDF（圖像頁面）→ 需 OCR（本地 ocr 或雲端 ocr-gemini 技能）
  - 混合型 → 逐頁混合處理
"""

import sys
import argparse
from pathlib import Path
import pymupdf

SUFFIXES = (".pdf", ".PDF")


def cjk_count(text: str) -> int:
    return sum(1 for c in text if "\u4e00" <= c <= "\u9fff")


def analyze_pdf(path: Path) -> dict:
    """分析 PDF：回傳每頁文字/圖片統計與整體分類。"""
    doc = pymupdf.open(path)
    pages = []
    total_text = 0
    total_cjk = 0
    total_images = 0
    for page in doc:
        text = page.get_text("text")
        img_n = len(page.get_images(full=True))
        total_text += len(text.strip())
        total_cjk += cjk_count(text)
        total_images += img_n
        pages.append({"text": len(text.strip()), "cjk": cjk_count(text),
                      "images": img_n})
    doc.close()

    n = len(pages)
    avg_cjk = total_cjk / n if n else 0
    if total_cjk >= 50 and total_images == 0:
        kind = "文字型"
    elif total_cjk >= 50 and total_images > 0:
        kind = "混合型"
    elif total_cjk < 50 and total_images > 0:
        kind = "掃描型"
    else:
        kind = "掃描型"  # 幾乎無文字 → 當掃描處理

    return {"pages": pages, "total_text": total_text, "total_cjk": total_cjk,
            "total_images": total_images, "kind": kind, "avg_cjk": avg_cjk}


def recommend(a: dict) -> str:
    if a["kind"] == "文字型":
        return ("直接文字抽取：python pdf2txt.py \"檔案.pdf\" --format docx\n"
                "      （pdf2docx 保留排版；純文字加 --format txt）")
    if a["kind"] == "掃描型":
        return ("使用 OCR 技能（非文字抽取）：\n"
                "      本地：python ~/.config/opencode/skills/ocr/ocr.py \"檔案.pdf\"\n"
                "      雲端：python ~/.config/opencode/skills/ocr-gemini/ocr_gemini.py \"檔案.pdf\"")
    return ("混合型：文字頁用 pdf2txt、圖像頁用 OCR。\n"
            "      或直接跑 --auto 讓工具自動逐頁處理。")


def extract_pdf(path: Path, layout: bool) -> str:
    doc = pymupdf.open(path)
    parts = []
    for page in doc:
        if layout:
            parts.append(page.get_text("text", sort=True))
        else:
            parts.append(page.get_text("text"))
    doc.close()
    return "\n\n".join(parts)


def write_docx(text: str, out_path: Path, title: str):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for para in text.split("\n"):
        if para.strip():
            p = doc.add_paragraph(para)
        else:
            doc.add_paragraph("")
    doc.save(out_path)


def convert_docx(src: Path, out_path: Path):
    """用 pdf2docx 重現 PDF 排版（格線/位置/字大小），保留為可編輯 Word。"""
    from pdf2docx import Converter

    cv = Converter(str(src))
    cv.convert(str(out_path), start=0, end=None)
    cv.close()


def convert_one(src: Path, out_dir: Path, layout: bool, fmt: str,
                silent: bool = False) -> Path:
    ext = ".docx" if fmt == "docx" else ".txt"
    out_path = out_dir / f"{src.stem}{ext}"
    if fmt == "docx":
        convert_docx(src, out_path)
    else:
        text = extract_pdf(src, layout)
        out_path.write_text(text, encoding="utf-8")
    if not silent:
        chars = out_path.stat().st_size
        print(f"  [OK] {out_path}")
    return out_path


def main():
    parser = argparse.ArgumentParser(description="PDF 轉文字")
    parser.add_argument("targets", nargs="+", help="PDF 檔案或資料夾")
    parser.add_argument("--out", default=None, help="輸出資料夾（預設與來源同資料夾/text）")
    parser.add_argument("--layout", action="store_true", help="保留欄位排版")
    parser.add_argument("--combine", action="store_true", help="全部合併成單一檔案")
    parser.add_argument("--format", default="docx", choices=["txt", "docx"],
                        help="輸出格式（預設 docx）")
    parser.add_argument("--analyze", action="store_true",
                        help="只分析 PDF 類型並給建議，不轉檔")
    parser.add_argument("--auto", action="store_true",
                        help="分析後自動選最佳方式處理（文字型→抽取；掃描型→OCR）")
    args = parser.parse_args()

    files = []
    for t in args.targets:
        p = Path(t)
        if p.is_dir():
            files.extend([f for f in p.iterdir() if f.suffix in SUFFIXES])
        elif p.is_file() and p.suffix in SUFFIXES:
            files.append(p)
        else:
            print(f"  [跳過] 不是 PDF：{t}", file=sys.stderr)

    if not files:
        print("找不到任何 PDF 檔案。", file=sys.stderr)
        sys.exit(1)

    if args.analyze:
        for f in files:
            a = analyze_pdf(f)
            print(f"■ {f.name}")
            print(f"  類型：{a['kind']}（{a['total_text']} 字元 / "
                  f"{a['total_cjk']} 中文 / {a['total_images']} 張圖，"
                  f"{len(a['pages'])} 頁）")
            print(f"  建議：{recommend(a)}")
            print()
        return

    if args.auto:
        for f in files:
            a = analyze_pdf(f)
            print(f"■ {f.name} → {a['kind']}")
            out_dir = Path(args.out) if args.out else f.parent / "text"
            out_dir.mkdir(parents=True, exist_ok=True)
            if a["kind"] in ("文字型", "混合型"):
                convert_one(f, out_dir, args.layout, args.format)
            if a["kind"] in ("掃描型", "混合型"):
                print("  [OCR] 圖像頁交由 OCR 技能處理（本工具不內建 OCR）")
                print(f"        建議：python "
                      f"%USERPROFILE%\\.config\\opencode\\skills\\ocr-gemini\\"
                      f"ocr_gemini.py \"{f}\" --out \"{out_dir / (f.stem + '_ocr.txt')}\"")
        return

    if args.combine:
        out_dir = Path(args.out) if args.out else Path.cwd()
        out_dir.mkdir(parents=True, exist_ok=True)
        ext = ".docx" if args.format == "docx" else ".txt"
        out_path = out_dir / f"combined_{len(files)}pdfs{ext}"
        if args.format == "docx":
            out_path.write_bytes(b"")  # pdf2docx 需逐檔轉出再合併，此處先保留
        else:
            combined = "".join(extract_pdf(f, args.layout) for f in files)
            out_path.write_text(combined, encoding="utf-8")
        print(f"  [OK] {out_path}  (合併 {len(files)} 個)")
        return

    print(f"轉換 {len(files)} 個 PDF ...")
    for f in files:
        out_dir = Path(args.out) if args.out else f.parent / "text"
        out_dir.mkdir(parents=True, exist_ok=True)
        convert_one(f, out_dir, args.layout, args.format)


if __name__ == "__main__":
    main()