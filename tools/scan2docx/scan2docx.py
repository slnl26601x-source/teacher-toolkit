"""
掃描書 PDF → 圖文並存的 Word（原圖 + Gemini OCR 文字）

用法：
  python scan2docx.py "掃描書.pdf"
  python scan2docx.py "掃描書.pdf" --pages 1-5            # 只處理指定頁
  python scan2docx.py "掃描書.pdf" --out "輸出.docx"       # 指定輸出檔名
  python scan2docx.py "掃描書.pdf" --model gemini-3.1-flash-lite
  python scan2docx.py "掃描書.pdf" --skip-ocr             # 只要原圖、不 OCR

逐頁智能判斷（Gemini 回報頁面類型）：
  - text_only（純文字頁）→ 不插圖，只留 OCR 文字
  - image_only（純圖無字）→ 保留原圖，不貼文字
  - mixed（圖文混合）→ 保留原圖 + OCR 文字

排版：
  - 連續排版：不做逐頁分頁，跨頁未完句子自動接合
  - 標題（章/節/目錄）用較大字體（16pt 黑體），正文 12pt 宋體、首行縮排
  - 孤行頁碼自動移除；段落依空行/縮排/對白引號重整

流程：
  1. 每頁抽成高解析 PNG
  2. Gemini 分析頁面類型 + OCR（快取於 _scan2docx_*/pages/*.ocr.txt）
  3. 組 Word：依類型決定插圖與文字

前置：~/.gemini.env 內有 GEMINI_API_KEY；需連網。
"""

import sys
import argparse
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")
sys.stderr.reconfigure(encoding="utf-8")


def parse_pages(spec: str, total: int):
    """解析 '1-5'、'3'、'1,3,5' 等頁碼規格，回傳 0-based 索引清單。"""
    if not spec:
        return list(range(total))
    idxs = []
    for part in spec.split(","):
        part = part.strip()
        if "-" in part:
            a, b = part.split("-", 1)
            a, b = int(a), int(b)
            idxs.extend(range(a - 1, b))
        else:
            idxs.append(int(part) - 1)
    return sorted({i for i in idxs if 0 <= i < total})


def extract_pages(pdf_path: Path, out_dir: Path, idxs, dpi: int = 200):
    import pymupdf
    doc = pymupdf.open(pdf_path)
    for i in idxs:
        pix = doc[i].get_pixmap(dpi=dpi)
        png = out_dir / f"page_{i + 1:03d}.png"
        pix.save(png)
    doc.close()


def _gemini_page(img_path: Path, model: str, prompt: str):
    """共用 Gemini 呼叫（組圖 + 重試 + JSON 解析）。回傳 (kind, text)。"""
    import base64
    import json
    from google import genai
    from google.genai import types

    key = None
    env = Path.home() / ".gemini.env"
    if env.exists():
        for line in env.read_text(encoding="utf-8").splitlines():
            if line.strip().startswith("GEMINI_API_KEY="):
                key = line.split("=", 1)[1].strip().strip('"').strip("'")
    if not key:
        print("錯誤：找不到 GEMINI_API_KEY（~/.gemini.env）", file=sys.stderr)
        sys.exit(1)

    client = genai.Client(api_key=key)
    image = types.Part.from_bytes(
        data=img_path.read_bytes(), mime_type="image/png")
    import time

    def _call():
        return client.models.generate_content(
            model=model,
            contents=[prompt, image],
        )

    resp = None
    last_wait = 0
    consecutive_429 = 0
    for attempt in range(15):
        try:
            resp = _call()
            break
        except Exception as e:
            msg = str(e)
            wait = 3
            if "429" in msg and "retryDelay" in msg:
                import re as _re
                m = _re.search(r"retryDelay': '(\d+)s", msg)
                if m:
                    wait = int(m.group(1))
            if "429" in msg:
                consecutive_429 += 1
            else:
                consecutive_429 = 0
            # 尊重 server 建議的 retryDelay，但封頂避免指數爆炸；連續 429 時才逐步放大
            if consecutive_429 <= 1:
                wait = min(wait, 30)
            else:
                wait = min(max(wait, 30), 90)
            last_wait = wait
            print(f"    429 限流，等待 {wait} 秒後重試（{attempt + 1}/15）...",
                  file=sys.stderr)
            time.sleep(wait)
    if resp is None:
        raise RuntimeError("Gemini OCR 多次重試仍失敗（限流）")
    raw = resp.text.strip() if resp.text else ""
    try:
        data = json.loads(raw.strip("` \n").removeprefix("json"))
    except Exception:
        # JSON 解析失敗：寬鬆抽取 "type" 與 "text"，避免把整包 JSON 當文字存檔
        import re as _re
        tm = _re.search(r'"type"\s*:\s*"(\w+)"', raw)
        kind2 = tm.group(1) if tm and tm.group(1) in (
            "text_only", "image_only", "mixed") else "mixed"
        text2 = ""
        im = _re.search(r'"text"\s*:\s*"((?:[^"\\]|\\.)*)"', raw, _re.S)
        if im:
            try:
                text2 = json.loads('"' + im.group(1) + '"')
            except Exception:
                text2 = im.group(1)
        # 若內容含未跳脫引號導致截斷，改從整段 JSON 抓取直到最後一個結尾引號
        if "}," not in text2 and text2 and not text2.rstrip().endswith(("。", "！", "？", "\"", "」", "』")):
            m2 = _re.search(r'"text"\s*:\s*"(.+)"\s*\}?\s*$', raw, _re.S)
            if m2:
                try:
                    text2 = json.loads('"' + m2.group(1) + '"')
                except Exception:
                    text2 = m2.group(1)
        data = {"type": kind2, "text": text2}
    kind = data.get("type") if data.get("type") in (
        "text_only", "image_only", "mixed") else "mixed"
    return kind, (data.get("text") or "").strip()


def ocr_page(img_path: Path, model: str):
    """標準 OCR：回傳 (kind, text)。kind: text_only | image_only | mixed"""
    return _gemini_page(img_path, model, _PROMPT_STD)


def ocr_page_formatted(img_path: Path, model: str):
    """關鍵頁格式版：回傳 (kind, text)。text 內段落可能帶格式標記：
    [章]章節大標題  [標]節/篇標題  [粗]粗體段  [中]置中段  [圖說]圖說。"""
    return _gemini_page(img_path, model, _PROMPT_FMT)


_PROMPT_STD = (
    "你是專業的影像分析助手。請分析這張圖片並回傳 JSON，格式如下：\n"
    '{"type": "text_only" 或 "image_only" 或 "mixed", "text": "..."}\n'
    "判斷規則（type 影響是否輸出原圖，務必依『內容主體』判斷）：\n"
    "- 頁面以文字為主（插圖/照片只是小範圍點綴或章首裝飾，正文仍佔大半）"
    "→ type=text_only，此時會輸出全部文字、不保留整頁圖。\n"
    "- 頁面幾乎只有插圖/照片、文字極少或只有圖說 → type=image_only。\n"
    "- 頁面文字與插圖/照片都佔相當比例、彼此並列 → type=mixed。\n"
    "text 欄位：用 OCR 完整提取圖片中的文字，嚴格保留原文語言（不翻譯），"
    "維持標點、段落與閱讀順序（直排中文依由上至下、由右至左）。\n"
    "文字整理要求：\n"
    "1. 依照原書的段落分塊：不同段落之間用一個空行（\\n\\n）分隔。\n"
    "2. 段落內部不要換行（把斷行接回同一段）。\n"
    "3. 段落輸出的先後順序：依閱讀順序排列——直排中文先由上至下讀完一欄，"
    "再由右欄往左欄；若頁面另有獨立圖說或底部文字，放在最後。\n"
    "4. 頁面頂部/底部的獨立頁碼（純數字的孤行）忽略，不要輸出。\n"
    "5. 若該頁是目錄/目次：條目後方或右側的頁碼數字「要保留」，"
    "直接接在該條目同一行末尾（例如『第一章　認識自然　15』）。\n"
    "6. 純數字行（孤行頁碼）不要輸出。\n"
    "7. 完整保留句子：若一句話在頁面上因圖片、留白或版面設計而被拆開、"
    "中斷或散落多處，請依上下文把它們補齊、接合成完整的一句，"
    "不要在中途截斷或重複。\n"
    "8. 圖片下方的圖說（說明文字）：若是獨立的一小段短句（約 80 字以內），"
    "請把它單獨一行、自成一段輸出，並在該段文字最前面加上『[圖說]』三個字（中括號），"
    "例如『[圖說]月光……有如流蕩的水光藻影。』，不要與相鄰正文合併。\n"
    "9. 若 type=mixed：每個段落開頭加上該段在頁面上的位置座標，格式為"
    "『(x,y)』，x 是水平位置（0-100，越大越靠右）、y 是垂直位置（0-100，越大越靠下），"
    "緊貼在段落文字最前面，例如『(62,35)月光……。』『(40,20)另一個笑說……。』。"
    "若 type=text_only 則不要加座標。\n"
    "image_only 時 text 留空字串。\n"
    "只回傳 JSON，不要加任何其他內容或 markdown 標記。"
)

_PROMPT_FMT = (
    "你是專業的影像分析助手。請分析這張圖片並回傳 JSON，格式如下：\n"
    '{"type": "text_only" 或 "image_only" 或 "mixed", "text": "..."}\n'
    "判斷規則（type 影響是否輸出原圖，務必依『內容主體』判斷）：\n"
    "- 頁面以文字為主（插圖/照片只是小範圍點綴或章首裝飾，正文仍佔大半）"
    "→ type=text_only，此時會輸出全部文字、不保留整頁圖。\n"
    "- 頁面幾乎只有插圖/照片、文字極少或只有圖說 → type=image_only。\n"
    "- 頁面文字與插圖/照片都佔相當比例、彼此並列 → type=mixed。\n"
    "text 欄位：用 OCR 完整提取圖片中的文字，嚴格保留原文語言（不翻譯），"
    "維持標點、段落與閱讀順序（直排中文依由上至下、由右至左）。\n"
    "文字整理要求：\n"
    "1. 依照原書的段落分塊：不同段落之間用一個空行（\\n\\n）分隔，"
    "段落內部不要換行。\n"
    "2. 每個段落，依它在原書的版面格式，在段落最前面加上一個標記（中括號）：\n"
    "   [章]  僅限真正的章節起始標題（含「第X章/篇」或「序/前言/後記/跋」等結構性標記），字級最大、通常置中\n"
    "   [標]  節/篇小標題、欄目標題（字級較大或粗體）\n"
    "   [粗]  粗體文字（整段或段中部分），包括頁首的獨立粗體短句"
    "（如「有这种女儿」「男生女生不象话」），不論是否含結構性標記\n"
    "   [中]  整段置中的文字（如詩、題辭、版權資訊、對白對話盒）\n"
    "   一般正文段落不加標記。段落若同時置中又大字 → 用 [章]；"
    "僅置中不粗不大字 → 用 [中]。\n"
    "   ⚠ 若一段中只有部分文字是粗體（如僅標題詞或開頭強調詞）"
    "而其餘不是，請拆成多個段落：粗體部分用 [粗]、非粗體部分不加標記。"
    "例如原文『有这种女儿老师的话：唉』中「有这种女儿」是粗體、"
    "其餘不是，則輸出：\\n[粗]有这种女儿\\n老师的话：唉...\n"
    "3. 頁面頂部/底部的獨立頁碼（純數字的孤行）忽略，不要輸出。\n"
    "4. 若該頁是目錄/目次：條目後方或右側的頁碼數字「要保留」，"
    "直接接在該條目同一行末尾（例如『第一章　認識自然　15』），"
    "目錄條目一律標 [標]（不要標 [章]，[章] 只給真正的章節起始標題頁）。\n"
    "5. 純數字行（孤行頁碼）不要輸出。\n"
    "6. 完整保留句子：若一句話在頁面上因圖片、留白或版面設計而被拆開、"
    "中斷或散落多處，請依上下文把它們補齊、接合成完整的一句。\n"
    "7. 圖片下方的圖說（說明文字）：獨立一小段短句（約 80 字以內）單獨一行、"
    "自成一段，並在該段最前面加上『[圖說]』，例如『[圖說]月光……』，"
    "不要與相鄰正文合併。\n"
    "8. 若 type=mixed：每個段落開頭加上該段在頁面上的位置座標『(x,y)』"
    "（x、y 皆 0-100），放在格式標記之後。若 type=text_only 則不加座標。\n"
    "image_only 時 text 留空字串。\n"
    "只回傳 JSON，不要加任何其他內容或 markdown 標記。"
)


_rapid_engine = None


def rapidocr_text(png: Path) -> str:
    """免費本機 OCR（RapidOCR，離線、無 API 額度）。失敗回傳空字串。"""
    global _rapid_engine
    try:
        if _rapid_engine is None:
            from rapidocr_onnxruntime import RapidOCR
            _rapid_engine = RapidOCR()
        result, _ = _rapid_engine(str(png))
    except Exception:
        return ""
    if not result:
        return ""
    lines = []
    for box, text, score in result:
        t = (text or "").strip()
        if t:
            lines.append(t)
    return "\n".join(lines)


def png_size(png: Path):
    """從 PNG 檔頭讀取寬高（無需額外套件）。失敗回傳 (None, None)。"""
    try:
        with open(png, "rb") as f:
            head = f.read(24)
        if head[:8] != b"\x89PNG\r\n\x1a\n":
            return None, None
        w = int.from_bytes(head[16:20], "big")
        h = int.from_bytes(head[20:24], "big")
        return w, h
    except Exception:
        return None, None


def rapidocr_layout(png: Path):
    """免費本機 OCR，回傳每行幾何資料（供排版判斷）。

    回傳 dict：
      {"W": 頁寬像素, "H": 頁高像素,
       "lines": [{"t": 文字, "x": 中心x, "y": 中心y,
                  "w": 寬, "h": 高}, ...]}   // x,y,w,h 皆歸一化 0-100
    失敗回傳 None。
    """
    global _rapid_engine
    try:
        if _rapid_engine is None:
            from rapidocr_onnxruntime import RapidOCR
            _rapid_engine = RapidOCR()
        result, _ = _rapid_engine(str(png))
    except Exception:
        return None
    if not result:
        return None
    W, H = png_size(png)
    if not W or not H:
        W, H = 1, 1
    out = []
    for box, text, score in result:
        t = (text or "").strip()
        if not t:
            continue
        xs = [p[0] for p in box]
        ys = [p[1] for p in box]
        w = max(xs) - min(xs)
        h = max(ys) - min(ys)
        cx = (min(xs) + max(xs)) / 2
        cy = (min(ys) + max(ys)) / 2
        out.append({
            "t": t,
            "x": round(cx * 100 / W, 2),
            "y": round(cy * 100 / H, 2),
            "w": round(w * 100 / W, 2),
            "h": round(h * 100 / H, 2),
        })
    return {"W": W, "H": H, "lines": out}


def analyze_title_lines(geo):
    """依幾何判斷標題行。回傳 list of (text, rank, centered)。

    rank: 2=章節大標題（字級明顯大） 1=標題級（置中或稍大） 0=一般
    centered: 該行水平置中（左右留白均等）

    穩健性：
    - 排除以數字為主的行（頁碼欄、條碼等被 RapidOCR 誤併成一長框）
    - 標題行需為短文字（避免長段被誤判）
    """
    import re
    import statistics
    if not geo or not geo.get("lines"):
        return []
    hs = [l["h"] for l in geo["lines"] if l.get("h", 0) > 0]
    if not hs:
        return []
    med = statistics.median(hs) or 1
    out = []
    for l in geo["lines"]:
        t = (l.get("t") or "").strip()
        if not t:
            continue
        size = l["h"] / med if med else 1
        # 排除數字為主的行（頁碼/條碼誤併長框）
        digits = sum(1 for ch in t if ch.isdigit())
        if digits and digits / len(t) >= 0.6 and len(t) >= 4:
            continue
        centered = abs(l["x"] + l["w"] / 2 - 50.0) < 6.0
        rank = 0
        if size >= 1.6 and len(t) <= 16:
            rank = 2 if (size >= 2.0 or (centered and size >= 1.5)) else 1
        elif centered and size >= 1.3 and len(t) <= 12:
            rank = 1
        out.append((t, rank, centered))
    return out


def clean_rapid_lines(lines):
    """過濾 RapidOCR 行：移除裸數字串（頁碼欄/條碼誤併）與過長異常行。

    括號包住的數字（如 (101)、（97））是合法內容（目錄頁碼/編號），保留。
    """
    import re
    out = []
    for t in lines:
        s = (t or "").strip()
        if not s:
            continue
        if re.fullmatch(r"[（(]?\s*\d{1,4}\s*[)）]?", s):
            # 純頁碼/編號行（可含括號）→ 保留（目錄頁碼行）
            out.append(s)
            continue
        digits = sum(1 for ch in s if ch.isdigit())
        if digits and digits / len(s) >= 0.6 and len(s) >= 4:
            continue
        if len(s) > 60:
            continue
        out.append(s)
    return out


def page_header_title(layout):
    """頁首孤立短標題行：第一行文字短（≤16 字）、位於頁面中上部、
    且與第二行間距明顯大於行高。典型文章標題頁（如『左手·右手』
    『五分钟跷课』），RapidOCR 字級中位數判斷不穩時仍可靠。
    回傳標題文字或 None。
    """
    lines = (layout or {}).get("lines") or []
    if len(lines) < 2:
        return None
    l0, l1 = lines[0], lines[1]
    t0 = (l0.get("t") or "").strip()
    if not t0 or len(t0) > 16:
        return None
    y0, h0 = l0.get("y", 0), l0.get("h", 0)
    if y0 < 8:          # 第一行貼頁頂 → 普通正文續頁
        return None
    hs = [l.get("h", 0) for l in lines if l.get("h", 0) > 0]
    med = sorted(hs)[len(hs) // 2] if hs else 0
    gap = l1.get("y", 0) - (y0 + h0)
    if gap >= 5 and gap >= 2.0 * (med or 1):
        return t0
    return None


def is_key_page(i, geo, title_marks, manual_pages=None):
    """是否為需要 Gemini 精準格式（粗體/版面）的關鍵頁。
    - 手動指定頁（--key-pages）
    - 全書前 3 頁（封面/版權/扉頁）
    - 含章節大標題（rank 2）的頁（章首頁）
    - 頁首孤立短標題行（文章標題頁）
    回傳原因字串或 None。
    """
    if manual_pages and i in manual_pages:
        return "手動指定"
    if i < 3:
        return "封面/前頁群"
    if any(m[1] == 2 for m in (title_marks or [])):
        return "章首大標題"
    if page_header_title(geo):
        return "文章標題頁"
    return None


_CJK_RE = r"[\u2e80-\u2fff\u3000-\u303f\u3400-\u4dbf\u4e00-\u9fff\uF900-\uFAFF\u3040-\u30ff\uFF01-\uFF60]"


def fullwidth_punct(s: str) -> str:
    """半形標點轉全形：相鄰為中文字（CJK）時轉換，位於英文詞內時保留。

    - ( ) → （ ）   , → ，   . → 。   ? → ？   ! → ！
    - : → ：   ; → ；   「『」等既為全形不變
    - 英文單字內的 - " ' .（如 Nokus-Ele、"Buffalo Bill"、St.、People's）維持半形
    """
    import re
    if not s or not re.search(r"[,()?!:;.']", s):
        return s
    full = {",": "，", "(": "（", ")": "）", "?": "？",
            "!": "！", ":": "：", ";": "；", ".": "。"}
    out = []
    for i, ch in enumerate(s):
        if ch not in full:
            out.append(ch)
            continue
        left = s[i - 1] if i > 0 else ""
        right = s[i + 1] if i + 1 < len(s) else ""
        # 兩側皆非中文（英文/數字/空白/符號）→ 視為英文語境，保留半形
        if not re.match(_CJK_RE, left) and not re.match(_CJK_RE, right):
            out.append(ch)
        else:
            out.append(full[ch])
    return "".join(out)


def reorder_mixed(text: str) -> str:
    """依直排閱讀順序（先由右至左、次由上至下）重排 mixed 頁段落。

    Gemini 會在 mixed 頁為每段標註 (x,y) 位置座標（0-100，x 越大越靠右、
    y 越大越靠下）。依「x 降序、y 升序」排序即符合「先右到左、次上到下」。
    無座標的段落（解析失敗）排在最後。
    """
    import re
    if not text:
        return text
    blocks = [b for b in text.split("\n\n") if b.strip()]
    items = []
    for b in blocks:
        # 可選格式標記前綴（如 [章]/[標]），後接半形或全形括號座標
        m = re.match(
            r"(\[[^\]]+\])?\s*[（(]\s*(\d+)\s*[,，]\s*(\d+)\s*[)）]\s*(.*)",
            b, re.S)
        if m:
            marker, x, y, body = m.group(1) or "", m.group(2), m.group(3), m.group(4)
            items.append((int(x), int(y), (marker + body).strip()))
        else:
            items.append((None, None, b.strip()))
    if not any(x is not None for x, _, _ in items):
        return text

    def key(it):
        x, y, _ = it
        return (-x, y) if x is not None else (9999, 0)

    ordered = sorted(items, key=key)
    # 座標剝離：只保留文字本體
    return "\n\n".join(t for _, _, t in ordered)


def reflow_text(text: str, toc: bool = False) -> list:
    """重整 OCR 文字為段落清單：
    - 移除孤行頁碼（如『8』，全行僅數字）
    - 空行、行首縮排、或行首為對話引號（「『）視為新段落
    - 其餘斷行直接接回同一段
    - toc=True（目錄頁）：每一行各自成一段，不整併
    """
    import re

    paras: list[str] = []
    cur: list[str] = []
    prev_was_toc = False
    for ln in text.splitlines():
        s = fullwidth_punct(ln.strip())
        if not s:
            if cur:
                paras.append("".join(cur))
                cur = []
            continue
        if toc:
            # 目錄頁：每行各自成段；但「純頁碼行」（(5)、12、（1））併入上一條目
            if re.fullmatch(r"[（(]?\s*\d{1,4}\s*[)）]?", s):
                if paras:
                    paras[-1] += s
                continue
            if cur:
                paras.append("".join(cur))
                cur = []
            paras.append(s)
            continue
        if re.fullmatch(r"\d{1,4}", s):  # 孤行頁碼
            continue
        toc_line = bool(re.match(r"^第[一二三四五六七八九十百千\d]+章\s*\S", s))
        starts_new = (
            s.startswith("　") or s.startswith("  ") or s.startswith("\t")
            or s.startswith("“") or s.startswith("「") or s.startswith("『")
            or s.startswith("[圖說]")   # Gemini 標記的圖說段
            or re.match(r"^\[[章標标粗中圖說]\]", s)  # Gemini 格式標記段（半/全形、繁簡）
            or bool(re.match(r"^[*＊●○■□・·]", s))   # 項次/註記標記分行
            or (toc_line and prev_was_toc)   # 目錄每章各成一行
        )
        if cur and starts_new:
            paras.append("".join(cur))
            cur = [s]
        else:
            cur.append(s)
        prev_was_toc = toc_line
    if cur:
        paras.append("".join(cur))
    return [p for p in paras if p.strip()]


def _merge_symbol_paras(paras: list) -> list:
    """將連續的單一符號段落（如 * * *、- - -）合併成一行，如『＊ ＊ ＊』"""
    import re
    out: list = []
    buf: list = []
    for p in paras:
        t = p.strip()
        if re.fullmatch(r"[*＊●○•·-]{1,3}", t) or re.fullmatch(r"[*＊●○•·-]", t):
            buf.append(t)
        else:
            if buf:
                out.append("  ".join(buf))
                buf = []
            out.append(p)
    if buf:
        out.append("  ".join(buf))
    return out


_CHAPTER_RE = r"^第[一二三四五六七八九十百千參貳\d]+[章篇編部]"
_CHAPTER_ONLY_RE = r"^第[一二三四五六七八九十百千參貳\d]+[章篇編部]\s*$"


def is_heading(s: str) -> bool:
    """判斷是否為標題（章/節/目錄等）"""
    import re
    t = s.strip()
    if not t:
        return False
    if re.match(_CHAPTER_RE, t):
        return True
    if re.match(r"^\d+[\.、．]\s*\S", t):
        return True
    norm = re.sub(r"\s", "", t)
    if norm in {"目录", "目次", "内容提要", "后记", "作者像", "序", "前言",
                "题记", "推薦序", "推荐序", "導言", "导言", "自序", "譯序",
                "譯後記", "跋"}:
        return True
    return False


def is_terminal(s: str) -> bool:
    """是否以句末標點（或閉引號）結束"""
    import re
    return bool(
        re.search(r"[。！？…\u2026」』”）)]\s*$", s)
        or re.search(r"[.!?]\s*$", s)
    )


def build_docx(pages_dir: Path, out_path: Path, idxs, model: str,
               skip_ocr: bool, dpi: int = 200, retry_model: str = "gemini-3.6-flash",
               retry_max: int = 5):
    import re
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "PMingLiU"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "PMingLiU")
    style.paragraph_format.first_line_indent = Pt(24)
    style.paragraph_format.space_after = Pt(6)

    items = []          # ("text"|"heading", 文字) 或 ("img", 路徑)
    prev_terminal = True   # 上一頁最後一段是否句末結束
    prev_has_img = False
    toc_region = False   # 是否在目錄區（逢「目次」頁進入，續頁沿用）
    prev_toc_page = False
    retry_used = [0]     # 強模型重試次數計數
    empty_pages = []     # 仍空白待補的頁碼

    for n, i in enumerate(idxs):
        png = pages_dir / f"page_{i + 1:03d}.png"
        if not png.exists():
            continue
        kind = "mixed"
        text = ""
        txt = pages_dir / f"page_{i + 1:03d}.ocr.txt"
        if txt.exists():
            try:
                kind, text = txt.read_text(encoding="utf-8").split("\n", 1)
            except Exception:
                kind, text = "mixed", ""
        elif not skip_ocr:
            kind, text = ocr_page(png, model)
            txt.write_text(f"{kind}\n{text}", encoding="utf-8")
        elif skip_ocr:
            continue   # 跳過 OCR 且無快取 → 略過該頁
        # 幾何分析（RapidOCR 存的 .geo.json）→ 標題/置中判斷
        geo = pages_dir / f"page_{i + 1:03d}.geo.json"
        title_marks = []
        if geo.exists():
            try:
                import json as _json
                layout = _json.loads(geo.read_text(encoding="utf-8"))
                title_marks = analyze_title_lines(layout)
            except Exception:
                title_marks = []
        if kind == "mixed" and (text or "").strip():
            # 依直排閱讀順序（右→左、上→下）重排段落
            text = reorder_mixed(text)
        if not skip_ocr:
            # 防呆：應有文字卻回傳空 → 先免費 RapidOCR，再限量強模型重試
            if kind in ("text_only", "mixed") and not (text or "").strip():
                rt = rapidocr_text(png)
                if rt.strip():
                    kind, text = "text_only", rt
                    txt.write_text(f"{kind}\n{text}", encoding="utf-8")
                    print(f"    頁 {i + 1} 文字為空，已用 RapidOCR 補回 {len(rt)} 字元",
                          file=sys.stderr)
                elif retry_used[0] < retry_max:
                    retry_used[0] += 1
                    print(f"    頁 {i + 1} RapidOCR 也空，改用 {retry_model} 重試 "
                          f"（{retry_used[0]}/{retry_max}）...", file=sys.stderr)
                    kind2, text2 = ocr_page(png, retry_model)
                    if text2.strip():
                        kind, text = kind2, text2
                        txt.write_text(f"{kind}\n{text}", encoding="utf-8")
                    else:
                        empty_pages.append(i + 1)
                else:
                    empty_pages.append(i + 1)
                    print(f"    頁 {i + 1} 仍空白（強模型額度已用罄，記錄待補）",
                          file=sys.stderr)

        page_has_img = kind != "text_only"
        page_img_pos = None
        page_captions = []          # 本頁圖說，插到圖片正下方
        # 頁首孤立短標題 → 分頁（非首頁、且前頁句末結束＝新文章起頁）
        if n > 0 and prev_terminal:
            try:
                hdr = page_header_title(layout) if geo.exists() and layout else None
                if hdr:
                    items.append(("page_break", None))
            except Exception:
                pass
        if page_has_img:
            items.append(("img", str(png)))
            page_img_pos = len(items) - 1

        def _is_toc_like(txt):
            """內容特徵偵測目錄：多數短行＋多個頁碼（括號頁碼或獨立數字行）。"""
            import re
            lines = [l.strip() for l in txt.splitlines() if l.strip()]
            if not lines:
                return False
            short = [l for l in lines if len(l) <= 40]
            numbered = [l for l in lines
                        if re.search(r"[（(]\s*\d+\s*[)）]$|^\d{1,3}$", l)]
            short_ratio = len(short) / len(lines)
            return (len(numbered) >= 5 and short_ratio >= 0.6)

        if "目次" in (text or ""):
            toc_region = True            # 進入目錄區（續頁沿用）
        elif toc_region and text:
            # 目錄續頁：整頁皆為短行 → 仍在目錄；出現長句 → 結束
            longest = max((len(l.strip()) for l in text.splitlines()), default=0)
            if longest > 45:
                toc_region = False
        elif text and _is_toc_like(text):
            toc_region = True            # 無「目次」字樣，但內容特徵像目錄
        is_toc_page = toc_region
        paras = (reflow_text(text, toc=is_toc_page)
                 if (text and kind != "image_only") else [])
        if not is_toc_page:
            paras = _merge_symbol_paras(paras)
        if is_toc_page:
            prev_terminal = True   # 目錄頁結束不跨頁接合
        if is_toc_page and not prev_toc_page:
            items.append(("toc_start", None))
        elif not is_toc_page and prev_toc_page:
            items.append(("toc_end", None))
        prev_toc_page = is_toc_page
        in_page_heading = False
        single_short_page = (
            len(paras) == 1 and 0 < len(paras[0].strip()) <= 18
            and not is_terminal(paras[0].strip())
        )
        for j, para in enumerate(paras):
            heading = is_heading(para) or (
                not is_toc_page and (
                    single_short_page
                    or (in_page_heading and len(para) <= 15 and not is_terminal(para))
                )
            )
            # 幾何標題（RapidOCR box 判斷：大字/置中）→ 補強 heading 判斷
            geo_rank = None
            if title_marks and not is_toc_page:
                p_norm = re.sub(r"\s", "", para)
                for t, rank, centered in title_marks:
                    t_norm = re.sub(r"\s", "", t)
                    if t_norm and (p_norm == t_norm or p_norm.startswith(t_norm)):
                        geo_rank = rank
                        break
                if geo_rank:
                    heading = True
            # Gemini 格式標記（關鍵頁）：[章]大標題 [標]標題 [粗]粗體 [中]置中
            # 相容半形/全形、繁簡體（標/标）
            fmt_mark = None
            m = re.match(r"\[([章標标粗中])\]", para)
            if m:
                fmt_mark = "[" + m.group(1) + "]"
                para = para[3:].strip()
                if not is_toc_page:
                    heading = True  # [章] 與 [標] 都是標題
            # 圖說（Gemini 標記的 [圖說] 段）→ 單獨成段、不與下文接合、剝離前綴
            caption = page_has_img and para.startswith("[圖說]")
            if caption:
                para = para[len("[圖說]"):].strip()
            # 章節標題頁（如「第一章」「第一篇」或幾何判為大標題）→ 分頁 + 置中
            # [章] 不含章節結構關鍵字時降級為 [粗]（如「有这种女儿」是粗體非章標題）
            if fmt_mark == "[章]" and not re.match(_CHAPTER_ONLY_RE, para) \
                    and not re.search(r"序|前言|後記|跋|自序|推薦序|導言", para):
                fmt_mark = "[粗]"
                heading = False
            # [標] 在續頁（前頁未句末結束）時降為正文（如「懷孕的女生」是續句非標題）
            if fmt_mark in ("[標]", "[标]") and not prev_terminal:
                fmt_mark = None
                heading = False
            is_chapter = (heading and re.match(_CHAPTER_ONLY_RE, para)) or (
                geo_rank == 2 or fmt_mark == "[章]"
            )
            if is_chapter:
                items.append(("page_break", None))
                items.append(("chapter_title", para))
                if j + 1 < len(paras):
                    sub = paras[j + 1].strip()
                    if (0 < len(sub) <= 18 and not is_terminal(sub)
                            and not re.match(r"^\[[章標粗中圖說]\]", sub)):
                        items[-1] = ("chapter_title", para + "  " + sub)
                        paras[j + 1] = ""   # 副標併入，避免重複輸出
                continue
            # 跨頁接合：上頁未句末結束 → 併入下頁首段（圖說不接合）
            if (
                j == 0 and not caption and not page_has_img and not prev_has_img
                and items and items[-1][0] in ("text", "heading")
                and not prev_terminal and not heading
            ):
                items[-1] = (items[-1][0], items[-1][1] + para)
            elif caption:
                page_captions.append(("caption", para))   # 圖說暫存，稍後插到圖後
            elif fmt_mark == "[粗]":
                items.append(("bold", para))              # 粗體正文段
            elif fmt_mark == "[中]":
                items.append(("center", para))            # 置中段
            else:
                items.append(("heading", para) if heading else ("text", para))
            if heading:
                in_page_heading = True
        # 圖說緊貼圖片下方（圖說說明的是那張圖，與正文區隔）
        if page_captions and page_img_pos is not None:
            for k, cap in enumerate(page_captions):
                items.insert(page_img_pos + 1 + k, cap)
        prev_terminal = True if (is_toc_page or single_short_page) else (
            (not paras) or is_terminal(paras[-1]))
        prev_has_img = page_has_img

    toc_mode = False   # 目前是否在目錄段落內
    for kind_, payload in items:
        if kind_ == "toc_start":
            toc_mode = True
            continue
        if kind_ == "toc_end":
            toc_mode = False
            continue
        if kind_ == "img":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(payload, width=Inches(4.5))
            continue
        if kind_ == "caption":
            # 圖說：緊貼圖下方、前綴「圖說：」、淺灰色小字、不縮排
            p = doc.add_paragraph()
            r = p.add_run("圖說：" + payload)
            r.font.size = Pt(10.5)
            r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            continue
        if kind_ == "page_break":
            doc.add_page_break()
            continue
        if kind_ == "chapter_title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(payload)
            r.bold = True
            r.font.size = Pt(22)
            r.font.name = "SimHei"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(72)
            p.paragraph_format.space_after = Pt(24)
            continue
        if kind_ == "bold":
            # 粗體正文段（Gemini 標記 [粗]）→ 粗體、其餘同正文
            p = doc.add_paragraph()
            r = p.add_run(payload)
            r.bold = True
            continue
        if kind_ == "center":
            # 置中段（Gemini 標記 [中]，如詩、題辭）→ 置中、不縮排
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(payload)
            p.paragraph_format.first_line_indent = Pt(0)
            continue

        norm = re.sub(r"\s", "", payload)
        if kind_ == "heading" and norm in ("目录", "目次"):
            toc_mode = True
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(payload)
            r.bold = True
            r.font.size = Pt(16)
            r.font.name = "SimHei"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_after = Pt(12)
            continue

        p = doc.add_paragraph()
        r = p.add_run(payload)
        toc_main = (
            re.match(_CHAPTER_RE, payload)
            or bool(re.match(r"^(推薦序|推荐序|導言|导言|自序|前言|後記|后记|題記|跋)", payload))
        )
        toc_num = bool(re.match(r"^\d+\s*\S", payload)) or bool(
            re.match(r"^[一二三四五六七八九十]+[、.．]", payload))
        in_toc_entry = toc_mode and (toc_main or toc_num or bool(
            re.match(_CHAPTER_RE, payload)))
        if kind_ == "heading" and not in_toc_entry:
            r.bold = True
            r.font.size = Pt(16)
            r.font.name = "SimHei"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            if toc_mode:
                # 目錄內的標題（如 推薦序/導言）→ 當主層級條目
                r.bold = True
                r.font.size = Pt(14)
                r.font.name = "SimHei"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.left_indent = Pt(0)
        else:
            if toc_mode or in_toc_entry:
                # 目錄條目：獨立行，依層級分字型大小
                p.paragraph_format.first_line_indent = Pt(0)
                if toc_main:
                    r.bold = True
                    r.font.size = Pt(14)
                    r.font.name = "SimHei"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
                    p.paragraph_format.left_indent = Pt(0)
                    p.paragraph_format.space_before = Pt(8)
                elif toc_num:
                    r.bold = True
                    r.font.size = Pt(12.5)
                    r.font.name = "PMingLiU"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "PMingLiU")
                    p.paragraph_format.left_indent = Pt(24)
                else:
                    r.font.size = Pt(12)
                    r.font.name = "PMingLiU"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "PMingLiU")
                    p.paragraph_format.left_indent = Pt(48)
    doc.save(out_path)
    if empty_pages:
        print(f"\n⚠ 以下頁面仍無文字（RapidOCR 與強模型均失敗），請手動補：{empty_pages}",
              file=sys.stderr)


def main():
    parser = argparse.ArgumentParser(description="掃描書 → 圖文 Word")
    parser.add_argument("pdf", help="掃描 PDF 路徑")
    parser.add_argument("--pages", default="", help="頁碼，如 1-5 或 1,3,5")
    parser.add_argument("--out", default=None, help="輸出 .docx 路徑")
    parser.add_argument("--model", default="gemini-3.1-flash-lite",
                        choices=["gemini-3.6-flash", "gemini-3.1-flash-lite",
                                 "gemini-2.5-pro"],
                        help="Gemini 模型")
    parser.add_argument("--skip-ocr", action="store_true", help="只要原圖")
    parser.add_argument("--dpi", type=int, default=200, help="抽圖解析度")
    parser.add_argument("--retry-model", default="gemini-3.6-flash",
                        choices=["gemini-3.6-flash", "gemini-3.1-flash-lite",
                                 "gemini-2.5-pro"],
                        help="空文字時的強模型（預設 gemini-3.6-flash）")
    parser.add_argument("--retry-max", type=int, default=5,
                        help="強模型每日/每次上限次數（預設 5）")
    parser.add_argument("--delay", type=float, default=4,
                        help="頁間等待秒數，避免觸發每分鐘上限（預設 4）")
    parser.add_argument("--no-gemini", action="store_true",
                        help="不使用 Gemini，全部用 RapidOCR（離線免費，無額度限制）")
    parser.add_argument("--key-pages", default="",
                        help="手動指定關鍵頁送 Gemini 精準格式（粗體/版面），如 1,3,5-7")
    args = parser.parse_args()

    pdf = Path(args.pdf)
    if not pdf.exists():
        print(f"錯誤：找不到 {pdf}", file=sys.stderr)
        sys.exit(1)

    import pymupdf
    total = len(pymupdf.open(pdf))
    idxs = parse_pages(args.pages, total)
    manual_pages = set(parse_pages(args.key_pages, total)) if args.key_pages else None
    if manual_pages:
        print(f"手動關鍵頁：{sorted(p + 1 for p in manual_pages)} 將送 Gemini 精準格式")

    work = pdf.parent / f"_scan2docx_{pdf.stem}"
    pages_dir = work / "pages"
    pages_dir.mkdir(parents=True, exist_ok=True)

    print(f"共 {total} 頁，處理 {len(idxs)} 頁（{args.pages or '全部'}）")
    print("步驟 1/3：抽取頁面圖 ...")
    extract_pages(pdf, pages_dir, idxs, args.dpi)

    if not args.skip_ocr:
        if args.no_gemini:
            print("步驟 2/3：RapidOCR 全頁離線辨識（不使用 Gemini）...")
            import time
            for i in idxs:
                png = pages_dir / f"page_{i + 1:03d}.png"
                txt = pages_dir / f"page_{i + 1:03d}.ocr.txt"
                geo = pages_dir / f"page_{i + 1:03d}.geo.json"
                if not txt.exists():
                    layout = rapidocr_layout(png)
                    rt = "\n".join(clean_rapid_lines(
                        [l["t"] for l in layout["lines"]])) if layout else ""
                    kind = "text_only" if rt.strip() else "image_only"
                    text = rt if kind == "text_only" else ""
                    txt.write_text(f"{kind}\n{text}", encoding="utf-8")
                    if layout:
                        geo.write_text(
                            __import__("json").dumps(layout, ensure_ascii=False),
                            encoding="utf-8")
                    if (i + 1) % 50 == 0:
                        print(f"  已處理 {i + 1}/{len(idxs)} 頁", file=sys.stderr)
                else:
                    kind = txt.read_text(encoding="utf-8").split("\n", 1)[0]
                print(f"  頁 {i + 1} ✓（{kind}）")
        else:
            print("步驟 2/3：混合 OCR（RapidOCR 幾何分析 → 關鍵頁 Gemini 精準格式）...")
            import time
            gemini_used = 0
            rapid_used = 0
            key_pages = []
            for i in idxs:
                png = pages_dir / f"page_{i + 1:03d}.png"
                txt = pages_dir / f"page_{i + 1:03d}.ocr.txt"
                geo = pages_dir / f"page_{i + 1:03d}.geo.json"
                if not txt.exists():
                    # 1) RapidOCR 幾何分析（免費）
                    layout = rapidocr_layout(png)
                    rt = "\n".join(clean_rapid_lines(
                        [l["t"] for l in layout["lines"]])) if layout else ""
                    title_marks = analyze_title_lines(layout) if layout else []
                    # 2) 關鍵頁（封面/目錄/章首）→ Gemini 精準格式 + 粗體
                    reason = is_key_page(i, layout, title_marks, manual_pages) if layout else (
                        "無文字待 Gemini" if not rt else None)
                    if reason and rt.strip():
                        kind, text = ocr_page_formatted(png, args.model)
                        gemini_used += 1
                        key_pages.append((i + 1, reason))
                        time.sleep(args.delay)
                    elif rt.strip():
                        kind, text = "text_only", rt
                        rapid_used += 1
                        if layout:
                            geo.write_text(
                                __import__("json").dumps(layout, ensure_ascii=False),
                                encoding="utf-8")
                    else:
                        kind, text = ocr_page(png, args.model)
                        gemini_used += 1
                        time.sleep(args.delay)
                    txt.write_text(f"{kind}\n{text}", encoding="utf-8")
                else:
                    kind = txt.read_text(encoding="utf-8").split("\n", 1)[0]
                print(f"  頁 {i + 1} ✓（{kind}）")
            print(f"  Gemini 呼叫: {gemini_used} 次、RapidOCR: {rapid_used} 次")
            for pg, rsn in key_pages:
                print(f"    - 頁 {pg} 送 Gemini（{rsn}）")
    else:
        print("步驟 2/3：跳過 OCR（--skip-ocr）")

    out = Path(args.out) if args.out else pdf.parent / f"{pdf.stem}_scan.docx"
    print(f"步驟 3/3：組 Word → {out}")
    build_docx(pages_dir, out, idxs, args.model, args.skip_ocr, args.dpi,
               args.retry_model, args.retry_max)
    print(f"完成：{out}")


if __name__ == "__main__":
    main()
